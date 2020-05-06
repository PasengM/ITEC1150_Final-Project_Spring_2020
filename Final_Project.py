################################## FINAL PROJECT #######################################################################
"""Program 1 : 15 points for code, 10 points for comments
   Program 2 : 40 points for code, 35 points for comments
   Extra-credit : 15 points
   Total : 100 points + 15 points ( as extra points ) """

################################## WHAT TO DO? #########################################################################
"""Program 1: The Taco Image (15 points for code, 10 points for comments) Search Unsplash for a taco image (for
example https://unsplash.com/photos/JiRSy0GfqPA) and save the image on your computer.

The downloaded image is very large. Use pillow to resize the image to a smaller size, perhaps no more than 800px wide
or tall (make sure you preserve the aspect ratio). Write the text "Random Taco Cookbook" on the image. Save the
modified image to a new file. """

################################## LIBRARIES ###########################################################################
# Imported libraries
from PIL import Image, ImageDraw, ImageFont  # From the Pillow library import Image (add image), ImageDraw (draw on
# image) and ImageFont (Write on it)
import requests  # Request = get json data from API servers
import docx  # Create my word document


################################# CODE PROGRAM 1: WORKING ON THE PICTURE ###############################################
# CODE PROGRAM 1
image = Image.open('Taco_spencer_davis_unsplash_Original.jpg')  # I select the picture I'll work with
size_max = (800, 800)  # Picture size that I want
image.thumbnail(size_max)  # I use .thumbnail to reduce the picture to the size I want. I also keep the ratio
image.save('800x800_taco_size.jpg')  # I save the picture

sized_picture = Image.open('800x800_taco_size.jpg')  # I select the picture I'll work with
image_draw = ImageDraw.Draw(sized_picture)  # I'll draw on the resized picture (800x800)
font = ImageFont.truetype('DejaVuSans.ttf', 45)  # I'll use the typology DejaVuSans.ttf as font + font size 45p
image_draw.text([40, 700], 'Random Taco Cookbook', fill='Fuchsia', font=font)  # I write 'Random Taco Cookbook'
# starting 40px to the right and 700px up to down

sized_picture.save('Image_Tacos_last_modified.jpg')  # I save this image with the name "Image_Tacos_last_modified.jpg"


################################## WHAT TO DO? #########################################################################
"""Program 2: Make the Random Taco Recipe Book (40 points for code, 35 points for comments) Use requests to download 
three random tacos from the random taco API. Save the data for each of three tacos in your program. Notice that each 
recipe is divided into five sections for base_layer, seasoning, mixin, condiment, and shell. Use Python to create a 
Word document. On the first page, insert the header text "Random Taco Cookbook" On the first page, add the resized 
taco image that you created with part 1. (hint: adding images to Word documents is covered in the textbook) On the 
first page, write the name of the image author On the first page, write the text of the random taco API URL On the 
first page, write your own name. On the second page, start writing the first taco recipe. Write all five components 
of the recipe. Use a larger font or heading for the heading for each of the sections. Please see example document for 
suggested style. After the first recipe, add a page break. To add another page, hint: google "python-docx add page 
break" Repeat to write all of the next recipe and a page break. Repeat to write all the third recipe. Save your word 
document. """


################################# CODE PROGRAM 2: ASSEMBLING THE BOOK RECIPES ##########################################

#######################     """Everything is on the first page """ #####################################################
# CODE PROGRAM 2

recipes_book = docx.Document()  # Create a new word document
recipes_book.add_paragraph('Random Taco Cookbook', 'Title')  # I add a new paragraph in this new document with
# 'Random Taco Cookbook' as header text and with title style
recipes_book.add_picture('Image_Tacos_last_modified.jpg', width=docx.shared.Inches(6),
                         height=docx.shared.Inches(6))  # I add the resized picture with 6 inches height and width
recipes_book.add_paragraph('Credits', 'Heading 1')  # Add the text "Credits" with the style "Heading 1" (subtitle) in
# a new paragraph

recipes_book.add_paragraph('Taco image: Photo by Spencer Davis on Unsplash', style='List Bullet')  # New paragraph
# with "Taco image..." with style bullet list
recipes_book.add_paragraph('Tacos from: http://taco-randomizer.herokuapp.com/random/?full_taco=true', style='List '
                                                                                                            'Bullet')
# # New paragraph with text "Taco from..." with style bullet list
recipes_book.add_paragraph('Code by: Paseng Moua', style='List Bullet')  # New paragraph with text "Code by: ..." with
# style bullet list

####################### """Everything is on the second,third, forth.... page """ #######################################

for item in range(3):  # for each item = it does it 3 times
    requests.get('https://taco-1150.herokuapp.com/random/?full_taco=true').json()  # I get the data from the json file
    recipe = requests.get('https://taco-1150.herokuapp.com/random/?full_taco=true').json()  # I store the data in a
    # variable "recipe"
    recipes_book.add_page_break()  # I add a new page
    recipes_book.add_paragraph(
        f'{recipe["base_layer"]["name"]} with {recipe["seasoning"]["name"]}, {recipe["condiment"]["name"]} and {recipe["mixin"]["name"]} in {recipe["shell"]["name"]}',
        'Title')  # I put the title of the recipe with all the elements asked : base_layer, seasoning, condiment,
    # mixin and shell as a 'Title'
    recipes_book.add_paragraph(f'{recipe["base_layer"]["name"]}', 'Heading 1')  # I add a new paragraph and put the
    # data in the recipe["base_layer"]["name"] as a subtitle 'Heading 1'
    recipes_book.add_paragraph(f'{recipe["base_layer"]["recipe"]}')  # I add a new paragraph and put the
    # data in the recipe["base_layer"]["recipe"] in the word document'
    recipes_book.add_paragraph(f'Seasoning: {recipe["seasoning"]["name"]}', 'Heading 1')  # I add a new paragraph and
    # put the data in the recipe["seasoning"]["name"] as a subtitle 'Heading 1'
    recipes_book.add_paragraph(f'{recipe["seasoning"]["recipe"]}')  # I add a new paragraph and put the
    # data in the recipe["seasoning"]["recipe"] in the word document'
    recipes_book.add_paragraph(f'Condiment: {recipe["condiment"]["name"]}', 'Heading 1')  # I add a new paragraph and
    # put the data in the recipe["condiment"]["name"] as a subtitle 'Heading 1'
    recipes_book.add_paragraph(f'{recipe["condiment"]["recipe"]}')  # I add a new paragraph and put the
    # data in the recipe["condiment"]["recipe"] in the word document'
    recipes_book.add_paragraph(f'Mixin: {recipe["mixin"]["name"]}', 'Heading 1')  # I add a new paragraph and
    # put the data in the recipe["mixin"]["name"] as a subtitle 'Heading 1'
    recipes_book.add_paragraph(f'{recipe["mixin"]["recipe"]}')  # I add a new paragraph and put the
    # data in the recipe["mixin"]["recipe"] in the word document'
    recipes_book.add_paragraph(f'Wrap: {recipe["shell"]["name"]}', 'Heading 1')  # I add a new paragraph and
    # put the data in the recipe["shell"]["name"] as a subtitle 'Heading 1'
    recipes_book.add_paragraph(f'{recipe["shell"]["recipe"]}')  # I add a new paragraph and put the
    # data in the recipe["shell"]["recipe"] in the word document'

recipes_book.save('Random Recipes Book.docx')  # I save the word document as 'Random Recipes Book.docx'


##################################  TODO_LIST   ########################################################################
# Try extra credits part
